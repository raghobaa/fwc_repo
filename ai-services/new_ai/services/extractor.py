import pdfplumber
import docx
import os


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error extracting PDF {file_path}: {e}")
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error extracting DOCX {file_path}: {e}")
    return text.strip()


def extract_text(file_path: str) -> str:
    """Auto-detect and extract text based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        return ""


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> list[str]:
    """Split text into overlapping chunks for better embedding coverage."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks
